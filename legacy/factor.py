#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
RVC Voice Similarity & Parameter Extractor – Professional Edition
- Word report with per‑figure numeric annotations and acoustic interpretations
- Multi‑metric speaker verification (Resemblyzer, F0, formants, spectrum, MFCC)
- Industry‑inspired decision logic with clear reasoning
- Save Word/docs to "report" folder, figures to "picture" folder
"""

import os, sys, glob, warnings, tempfile, textwrap, subprocess
import numpy as np
import parselmouth
import soundfile as sf
import sounddevice as sd
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import matplotlib
matplotlib.use("TkAgg")
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
import librosa, librosa.display

# Unified F0 extraction parameters (fix #2: keep consistent across all calls)
PITCH_TIME_STEP = 0.01
PITCH_FLOOR = 65       # Hz
PITCH_CEILING = 2093   # Hz

# ---------- Optional imports ----------
try:
    from resemblyzer import VoiceEncoder, preprocess_wav
    RESEMBLYZER_AVAILABLE = True
except ImportError:
    RESEMBLYZER_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

# Cached speaker encoder (fix #5: load model only once across the app lifetime)
_ENCODER = None
def _get_encoder():
    global _ENCODER
    if RESEMBLYZER_AVAILABLE and _ENCODER is None:
        _ENCODER = VoiceEncoder()
    return _ENCODER

warnings.filterwarnings("ignore")

# ==================== Audio Feature Extraction ====================
def extract_pitch_array(audio, sr=16000):
    snd = parselmouth.Sound(audio, sampling_frequency=sr)
    pitch = snd.to_pitch(time_step=PITCH_TIME_STEP, pitch_floor=PITCH_FLOOR, pitch_ceiling=PITCH_CEILING)
    f0 = pitch.selected_array['frequency']
    valid = f0[f0 > 0]
    if len(valid) == 0:
        raise ValueError("No voiced frames detected")
    median_pitch = np.median(valid)
    gender = np.clip((median_pitch - 120) / (210 - 120), 0.0, 1.0)
    return f0, median_pitch, gender

def extract_formants(audio, sr=16000):
    snd = parselmouth.Sound(audio, sampling_frequency=sr)
    formant = snd.to_formant_burg(time_step=0.01, max_number_of_formants=5,
                                  maximum_formant=5500, window_length=0.025)
    times = formant.xs()
    f1 = [formant.get_value_at_time(1, t) for t in times]
    f2 = [formant.get_value_at_time(2, t) for t in times]
    f3 = [formant.get_value_at_time(3, t) for t in times]
    return np.array(times), np.array(f1), np.array(f2), np.array(f3)

def get_audio_features(audio, sr):
    dur = len(audio) / sr
    rms = np.sqrt(np.mean(audio**2))
    zcr = np.mean(librosa.feature.zero_crossing_rate(y=audio))
    cent = np.mean(librosa.feature.spectral_centroid(y=audio, sr=sr))
    return {
        'duration': round(dur, 2),
        'rms': round(float(rms), 4),
        'zcr': round(float(zcr), 4),
        'centroid': round(float(cent), 1)
    }

# ==================== Similarity Metrics ====================
def compute_pitch_similarity(p1, p2):
    semis = abs(12 * np.log2(p2 / p1))
    return float(np.clip(1 - semis / 12.0, 0.0, 1.0))

def compute_mfcc_similarity(a1, sr1, a2, sr2):
    try:
        m1 = librosa.feature.mfcc(y=a1.astype(np.float64), sr=sr1, n_mfcc=13)
        m2 = librosa.feature.mfcc(y=a2.astype(np.float64), sr=sr2, n_mfcc=13)
        v1, v2 = np.mean(m1, axis=1), np.mean(m2, axis=1)
        dot = np.dot(v1, v2)
        norm = np.linalg.norm(v1) * np.linalg.norm(v2)
        return float(np.clip(dot / norm, -1, 1)) if norm > 0 else 0.0
    except Exception:
        return 0.0

def compute_speaker_similarity(a1, sr1, a2, sr2):
    if RESEMBLYZER_AVAILABLE:
        try:
            enc = _get_encoder()  # fix #5: reuse cached encoder instead of reloading
            w1 = preprocess_wav(a1.astype(np.float64), source_sr=sr1)
            w2 = preprocess_wav(a2.astype(np.float64), source_sr=sr2)
            e1, e2 = enc.embed_utterance(w1), enc.embed_utterance(w2)
            return float(np.clip((np.dot(e1, e2) + 1) / 2.0, 0.0, 1.0))
        except Exception as e:
            print(f"Resemblyzer failed: {e}")
    mfcc_sim = compute_mfcc_similarity(a1, sr1, a2, sr2)
    return float(np.clip((mfcc_sim + 1) / 2.0, 0.0, 1.0))

# ==================== Multi‑metric Decision Engine ====================
def analyze_speaker_verification(src_audio, src_sr, tgt_audio, tgt_sr):
    """
    Returns a dict with all relevant metrics and a final decision.
    Combines:
      - Speaker embedding similarity (Resemblyzer)
      - Pitch similarity
      - Formant F2 mean difference (normalized)
      - Spectral centroid difference (normalized)
      - MFCC cosine similarity (global timbre)
    """
    # 1. Speaker embedding (0-1)
    spk_sim = compute_speaker_similarity(src_audio, src_sr, tgt_audio, tgt_sr)

    # 2. Pitch
    _, src_pitch, _ = extract_pitch_array(src_audio, src_sr)
    _, tgt_pitch, _ = extract_pitch_array(tgt_audio, tgt_sr)
    pitch_sim = compute_pitch_similarity(src_pitch, tgt_pitch)
    semitones = 12 * np.log2(tgt_pitch / src_pitch)

    # 3. Formant F2 means
    _, _, f2_src, _ = extract_formants(src_audio, src_sr)
    _, _, f2_tgt, _ = extract_formants(tgt_audio, tgt_sr)
    f2_diff = abs(np.nanmean(f2_src) - np.nanmean(f2_tgt))
    # Normalize: typical speaker variability ~ 200-300 Hz, use 300 Hz as reference
    f2_norm = np.clip(f2_diff / 300.0, 0.0, 1.0)

    # 4. Spectral centroid
    src_feat = get_audio_features(src_audio, src_sr)
    tgt_feat = get_audio_features(tgt_audio, tgt_sr)
    cent_diff = abs(src_feat['centroid'] - tgt_feat['centroid'])
    cent_norm = np.clip(cent_diff / 1000.0, 0.0, 1.0)

    # 5. MFCC cosine similarity (global timbre)
    mfcc_sim = compute_mfcc_similarity(src_audio, src_sr, tgt_audio, tgt_sr)
    mfcc_sim_01 = float(np.clip((mfcc_sim + 1) / 2.0, 0.0, 1.0))

    # Combine into weighted decision score (industry inspired)
    # Weights: speaker embedding 0.4, pitch 0.2, formant 0.2, centroid 0.1, mfcc 0.1
    decision_score = (0.4 * spk_sim +
                      0.2 * pitch_sim +
                      0.2 * (1 - f2_norm) +
                      0.1 * (1 - cent_norm) +
                      0.1 * mfcc_sim_01)

    # Final decision
    if spk_sim > 0.85 and decision_score > 0.78:
        verdict, confidence = "YES", "High"
    elif spk_sim > 0.78 and decision_score > 0.68 and pitch_sim > 0.6 and f2_norm < 0.5:
        verdict, confidence = "YES", "Medium"
    elif spk_sim > 0.70 and decision_score > 0.60 and pitch_sim > 0.5:
        verdict, confidence = "Uncertain", "Low"
    else:
        verdict, confidence = "NO", "Very Low / Different"

    # Build reasoning text
    reasons = []
    if spk_sim > 0.82:
        reasons.append("Speaker embedding strongly similar (timbre)")
    elif spk_sim > 0.70:
        reasons.append("Speaker embedding moderately similar")
    else:
        reasons.append("Speaker embedding significantly different")

    if pitch_sim > 0.75:
        reasons.append("pitch range closely matches")
    elif pitch_sim > 0.5:
        reasons.append(f"pitch differs by {abs(semitones):.1f} semitones (moderate)")
    else:
        reasons.append(f"pitch differs by {abs(semitones):.1f} semitones (large)")

    if f2_norm < 0.3:
        reasons.append("formant F2 means very close (similar vocal tract)")
    elif f2_norm < 0.6:
        reasons.append(f"formant F2 differs by {f2_diff:.0f} Hz (noticeable)")
    else:
        reasons.append(f"formant F2 differs by {f2_diff:.0f} Hz (strong difference)")

    if cent_norm < 0.3:
        reasons.append("spectral centroid very close")
    elif cent_norm < 0.6:
        reasons.append("spectral centroid moderately different")
    else:
        reasons.append("spectral centroid clearly different")

    full_reason = ". ".join(reasons) + "."
    if verdict == "YES":
        full_reason += " Multiple acoustic cues align, supporting same‑speaker hypothesis."
    elif verdict == "NO":
        full_reason += " Multiple acoustic cues diverge, indicating different speakers."
    else:
        full_reason += " Evidence is mixed; further manual inspection recommended."

    return {
        'speaker_sim': spk_sim,
        'pitch_sim': pitch_sim,
        'f2_diff': f2_diff,
        'f2_norm': f2_norm,
        'cent_diff': cent_diff,
        'cent_norm': cent_norm,
        'mfcc_sim': mfcc_sim_01,
        'decision_score': decision_score,
        'verdict': verdict,
        'confidence': confidence,
        'reason': full_reason,
        'src_pitch': src_pitch,
        'tgt_pitch': tgt_pitch,
        'semitones': semitones,
        'src_centroid': src_feat['centroid'],
        'tgt_centroid': tgt_feat['centroid']
    }

# ==================== Recorder ====================
class Recorder:
    def __init__(self, sr=16000):
        self.sr = sr
        self.recording = False
        self.frames = []
        self.stream = None

    def start(self):
        self.frames = []
        self.recording = True
        self.stream = sd.InputStream(samplerate=self.sr, channels=1, dtype='float32', callback=self._cb)
        self.stream.start()

    def _cb(self, indata, frames, time, status):
        if self.recording:
            self.frames.append(indata.copy())

    def stop(self):
        self.recording = False
        if self.stream:
            self.stream.stop()
            self.stream.close()
            self.stream = None
        if not self.frames:
            return None
        return np.concatenate(self.frames, axis=0).flatten().astype(np.float64)

# ==================== Plotting (enhanced with annotations) ====================
def annotate_text(ax, text, x=0.02, y=0.95, fontsize=8):
    ax.text(x, y, str(text), transform=ax.transAxes, fontsize=fontsize,
            verticalalignment='top', bbox=dict(boxstyle='round', facecolor='white', alpha=0.85))

def plot_waveform(ax, audio, sr, title, color):
    t = np.linspace(0, len(audio)/sr, len(audio))
    ax.plot(t, audio, color=color, linewidth=0.7)
    ax.set_title(title, fontsize=10)
    ax.set_xlabel("Time (s)")

def plot_spectrogram(ax, audio, sr, title):
    D = librosa.stft(audio.astype(np.float64))
    S_db = librosa.amplitude_to_db(np.abs(D), ref=np.max)
    librosa.display.specshow(S_db, sr=sr, x_axis='time', y_axis='hz', ax=ax, cmap='inferno')
    ax.set_title(title, fontsize=10)
    ax.set_ylim(0, 3000)

def plot_f0_contour(ax, times, f0, title, color):
    ax.plot(times, f0, color=color, linewidth=1)
    ax.set_title(title, fontsize=10)
    ax.set_xlabel("Time (s)")
    ax.set_ylabel("F0 (Hz)")

def plot_formants(ax, times, f1, f2, f3, title):
    ax.plot(times, f1, 'r', label='F1')
    ax.plot(times, f2, 'g', label='F2')
    ax.plot(times, f3, 'b', label='F3')
    ax.set_title(title, fontsize=10)
    ax.set_xlabel("Time (s)")
    ax.set_ylabel("Frequency (Hz)")
    ax.legend(fontsize=8)

def plot_avg_spectrum(ax, audio, sr, label, color):
    D = np.abs(librosa.stft(audio.astype(np.float64)))
    avg = np.mean(D, axis=1)
    freqs = librosa.fft_frequencies(sr=sr)
    db = librosa.amplitude_to_db(avg, ref=np.max)
    ax.semilogx(freqs, db, color=color, label=label, linewidth=1)
    ax.set_xlim(50, sr//2)
    ax.set_xlabel("Frequency (Hz)")
    ax.set_ylabel("dB")
    ax.legend(fontsize=8)
    ax.grid(True, alpha=0.3)

def plot_mfcc_heatmap(ax, audio, sr, title):
    mfcc = librosa.feature.mfcc(y=audio.astype(np.float64), sr=sr, n_mfcc=13)
    librosa.display.specshow(mfcc, sr=sr, x_axis='time', ax=ax, cmap='coolwarm')
    ax.set_title(title, fontsize=10)

def generate_report_figures(src_audio, src_sr, tgt_audio, tgt_sr,
                            src_id, tgt_id, metrics, decision,
                            output_dir):
    """
    Save all figures to output_dir (will be created if needed).
    Returns list of saved image paths.
    """
    os.makedirs(output_dir, exist_ok=True)
    img_paths = []

    src_dur = len(src_audio) / src_sr
    tgt_dur = len(tgt_audio) / tgt_sr
    common_dur = min(src_dur, tgt_dur)

    src_aligned = src_audio[:int(common_dur * src_sr)] if src_dur > common_dur else src_audio
    tgt_aligned = tgt_audio[:int(common_dur * tgt_sr)] if tgt_dur > common_dur else tgt_audio

    # ---- Fig 1: Overview ----
    fig1, axes = plt.subplots(3, 2, figsize=(10, 7))
    fig1.suptitle(f"Overview: {src_id} vs {tgt_id} (first {common_dur:.1f}s)", fontweight='bold')
    plot_waveform(axes[0,0], src_aligned, src_sr, f"{src_id} - Waveform", 'steelblue')
    plot_waveform(axes[0,1], tgt_aligned, tgt_sr, f"{tgt_id} - Waveform", 'darkorange')
    plot_spectrogram(axes[1,0], src_aligned, src_sr, f"{src_id} - Spectrogram")
    plot_spectrogram(axes[1,1], tgt_aligned, tgt_sr, f"{tgt_id} - Spectrogram")

    src_f0, src_pitch, _ = extract_pitch_array(src_audio, src_sr)
    tgt_f0, tgt_pitch, _ = extract_pitch_array(tgt_audio, tgt_sr)
    axes[2,0].hist(src_f0[src_f0>0], bins=30, color='steelblue', alpha=0.7)
    axes[2,0].axvline(src_pitch, color='red', linestyle='--', label=f"Median {src_pitch:.1f} Hz")
    axes[2,0].set_title(f"{src_id} - F0 Distribution", fontsize=10)
    axes[2,0].legend(fontsize=8)
    axes[2,1].hist(tgt_f0[tgt_f0>0], bins=30, color='darkorange', alpha=0.7)
    axes[2,1].axvline(tgt_pitch, color='red', linestyle='--', label=f"Median {tgt_pitch:.1f} Hz")
    axes[2,1].set_title(f"{tgt_id} - F0 Distribution", fontsize=10)
    axes[2,1].legend(fontsize=8)
    plt.tight_layout()
    p1 = os.path.join(output_dir, "overview.png")
    fig1.savefig(p1, dpi=150)
    plt.close(fig1)
    img_paths.append(p1)

    # ---- Fig 2: F0 contours ----
    snd_src = parselmouth.Sound(src_aligned, sampling_frequency=src_sr)
    pitch_src = snd_src.to_pitch(time_step=PITCH_TIME_STEP, pitch_floor=PITCH_FLOOR, pitch_ceiling=PITCH_CEILING)
    times_src = pitch_src.xs()
    f0_src = pitch_src.selected_array['frequency']
    snd_tgt = parselmouth.Sound(tgt_aligned, sampling_frequency=tgt_sr)
    pitch_tgt = snd_tgt.to_pitch(time_step=PITCH_TIME_STEP, pitch_floor=PITCH_FLOOR, pitch_ceiling=PITCH_CEILING)
    times_tgt = pitch_tgt.xs()
    f0_tgt = pitch_tgt.selected_array['frequency']
    fig2, (ax1, ax2) = plt.subplots(2, 1, figsize=(8, 4))
    plot_f0_contour(ax1, times_src, f0_src, f"{src_id} - F0 Contour", 'steelblue')
    annotate_text(ax1, f"Median: {src_pitch:.1f} Hz")
    plot_f0_contour(ax2, times_tgt, f0_tgt, f"{tgt_id} - F0 Contour", 'darkorange')
    annotate_text(ax2, f"Median: {tgt_pitch:.1f} Hz")
    plt.tight_layout()
    p2 = os.path.join(output_dir, "f0_contours.png")
    fig2.savefig(p2, dpi=150)
    plt.close(fig2)
    img_paths.append(p2)

    # ---- Fig 3: Formants ----
    t_src, f1_src, f2_src, f3_src = extract_formants(src_aligned, src_sr)
    t_tgt, f1_tgt, f2_tgt, f3_tgt = extract_formants(tgt_aligned, tgt_sr)
    f1_mean_src, f2_mean_src = float(np.nanmean(f1_src)), float(np.nanmean(f2_src))
    f1_mean_tgt, f2_mean_tgt = float(np.nanmean(f1_tgt)), float(np.nanmean(f2_tgt))
    fig3, (ax1, ax2) = plt.subplots(2, 1, figsize=(8, 4))
    plot_formants(ax1, t_src, f1_src, f2_src, f3_src, f"{src_id} - Formants")
    annotate_text(ax1, f"Mean F1={f1_mean_src:.0f} Hz, F2={f2_mean_src:.0f} Hz")
    plot_formants(ax2, t_tgt, f1_tgt, f2_tgt, f3_tgt, f"{tgt_id} - Formants")
    annotate_text(ax2, f"Mean F1={f1_mean_tgt:.0f} Hz, F2={f2_mean_tgt:.0f} Hz")
    plt.tight_layout()
    p3 = os.path.join(output_dir, "formants.png")
    fig3.savefig(p3, dpi=150)
    plt.close(fig3)
    img_paths.append(p3)

    # ---- Fig 4: Average spectrum ----
    fig4, ax = plt.subplots(figsize=(6, 4))
    plot_avg_spectrum(ax, src_audio, src_sr, src_id, 'steelblue')
    plot_avg_spectrum(ax, tgt_audio, tgt_sr, tgt_id, 'darkorange')
    ax.set_title("Average Spectrum (full duration)")
    src_cent = get_audio_features(src_audio, src_sr)['centroid']
    tgt_cent = get_audio_features(tgt_audio, tgt_sr)['centroid']
    annotate_text(ax, f"Centroid: Src={src_cent} Hz, Tgt={tgt_cent} Hz")
    plt.tight_layout()
    p4 = os.path.join(output_dir, "avg_spectrum.png")
    fig4.savefig(p4, dpi=150)
    plt.close(fig4)
    img_paths.append(p4)

    # ---- Fig 5: MFCC ----
    fig5, (ax1, ax2) = plt.subplots(2, 1, figsize=(6, 5))
    plot_mfcc_heatmap(ax1, src_aligned, src_sr, f"{src_id} - MFCC")
    plot_mfcc_heatmap(ax2, tgt_aligned, tgt_sr, f"{tgt_id} - MFCC")
    plt.tight_layout()
    p5 = os.path.join(output_dir, "mfcc.png")
    fig5.savefig(p5, dpi=150)
    plt.close(fig5)
    img_paths.append(p5)

    return img_paths

def combine_images_vertically(img_paths, output_path):
    if not PIL_AVAILABLE: return False
    images = [Image.open(p) for p in img_paths]
    widths, heights = zip(*(i.size for i in images))
    max_width = max(widths)
    total_height = sum(heights)
    combined = Image.new('RGB', (max_width, total_height), 'white')
    y_offset = 0
    for im in images:
        combined.paste(im, (0, y_offset))
        y_offset += im.size[1]
    combined.save(output_path)
    return True

def create_word_report(img_paths, src_id, tgt_id, metrics, decision, output_docx):
    if not DOCX_AVAILABLE: return False
    doc = Document()
    doc.styles['Normal'].font.size = Pt(10)

    doc.add_heading(f'Voice Similarity Report: {src_id} vs {tgt_id}', 0)

    # Summary
    doc.add_paragraph(f"Conclusion: {decision['verdict']} (confidence: {decision['confidence']})", style='Intense Quote')
    doc.add_paragraph(f"Reasoning: {decision['reason']}")

    # Metrics table
    doc.add_heading('Quantitative Metrics', level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = 'Metric'
    hdr[1].text = 'Value'
    for key, val in metrics.items():
        row = table.add_row().cells
        row[0].text = key
        row[1].text = str(val)

    # Detailed per‑figure explanations
    doc.add_heading('Figure Explanations', level=2)

    explanations = [
        ("Figure 1: Overview",
         "The first row shows waveform amplitude envelopes. "
         "The second row displays spectrograms (0‑3000 Hz); similar harmonic structure and formant patterns suggest similar vocal tract. "
         "The third row gives F0 histograms with median pitch (red dashed line). "
         "Overlapping distributions indicate similar pitch range."),
        ("Figure 2: F0 Contours",
         "Pitch tracks over time for the aligned segment. "
         "Similar contour shapes (e.g., declination, intonation patterns) are typical of the same speaker or similar speaking style."),
        ("Figure 3: Formant Tracks",
         "Formants F1, F2, F3 reflect vocal tract resonances. "
         "Mean F1 and F2 are strong indicators of speaker identity; differences >15% often point to different speakers."),
        ("Figure 4: Average Spectrum",
         "Long‑term average spectrum shows overall spectral slope and brightness. "
         "Spectral centroid (annotated) quantifies timbre brightness; large differences (>500 Hz) suggest different recording conditions or speakers."),
        ("Figure 5: MFCC Heatmaps",
         "Mel‑frequency cepstral coefficients capture timbral detail. "
         "Visual similarity in the heatmap patterns, together with the MFCC cosine similarity score, reflects global voice similarity.")
    ]

    for i, (title, desc) in enumerate(explanations, 1):
        doc.add_heading(title, level=3)
        doc.add_paragraph(desc)
        doc.add_picture(img_paths[i-1], width=Inches(5.5))
        doc.add_paragraph()  # spacing

    doc.save(output_docx)
    return True

def open_file_with_default_app(filepath):
    if sys.platform == "win32":
        os.startfile(filepath)
    elif sys.platform == "darwin":
        subprocess.run(["open", filepath], check=False)  # fix #4: arg list, no shell
    else:
        subprocess.run(["xdg-open", filepath], check=False)

# ==================== GUI ====================
class PitchGUI:
    def __init__(self, root):
        self.root = root
        root.title("Voice Similarity & RVC Extractor")
        root.geometry("920x720")
        root.minsize(860, 640)

        self.src_audio = None
        self.src_sr = 16000
        self.src_recorder = Recorder(sr=16000)

        self.tgt_audio = None
        self.tgt_sr = 16000
        self.tgt_recorder = Recorder(sr=16000)

        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        output_dir = os.path.join(base_dir, "output")
        self.audio_dir = os.path.join(output_dir, "audio")
        self.report_dir = os.path.join(output_dir, "report")
        self.picture_dir = os.path.join(output_dir, "picture")
        os.makedirs(self.audio_dir, exist_ok=True)
        os.makedirs(self.report_dir, exist_ok=True)
        os.makedirs(self.picture_dir, exist_ok=True)

        self._apply_theme()
        self._build_ui()

    def _apply_theme(self):
        self.ui = {
            "bg": "#0B1020",
            "panel": "#0F1730",
            "panel_2": "#111C3A",
            "border": "#22305A",
            "text": "#EAF0FF",
            "muted": "#9FB0D5",
            "primary": "#6D5EF6",
            "primary_hover": "#5A4BE8",
            "success": "#22C55E",
            "danger": "#EF4444",
            "warn": "#F59E0B",
            "focus": "#93C5FD",
            "mono": ("Cascadia Mono", 10),
            "sans": ("Segoe UI", 10),
            "sans_bold": ("Segoe UI Semibold", 10),
            "title": ("Segoe UI Semibold", 14),
        }

        self.root.configure(bg=self.ui["bg"])
        style = ttk.Style(self.root)
        try:
            style.theme_use("clam")
        except tk.TclError:
            pass

        style.configure(".", font=self.ui["sans"])
        style.configure("App.TFrame", background=self.ui["bg"])
        style.configure("Panel.TFrame", background=self.ui["panel"], borderwidth=1, relief="solid")
        style.configure("Card.TFrame", background=self.ui["panel_2"], borderwidth=1, relief="solid")
        style.configure("Title.TLabel", background=self.ui["bg"], foreground=self.ui["text"], font=self.ui["title"])
        style.configure("Sub.TLabel", background=self.ui["bg"], foreground=self.ui["muted"])
        style.configure("Field.TLabel", background=self.ui["panel_2"], foreground=self.ui["muted"])
        style.configure("Value.TLabel", background=self.ui["panel_2"], foreground=self.ui["text"], font=self.ui["sans_bold"])

        style.configure(
            "Modern.TEntry",
            fieldbackground=self.ui["panel"],
            background=self.ui["panel"],
            foreground=self.ui["text"],
            bordercolor=self.ui["border"],
            lightcolor=self.ui["border"],
            darkcolor=self.ui["border"],
            insertcolor=self.ui["text"],
        )
        style.map("Modern.TEntry", bordercolor=[("focus", self.ui["focus"])])

        style.configure("Modern.TButton", padding=(12, 8), background=self.ui["panel"], foreground=self.ui["text"])
        style.map("Modern.TButton", background=[("active", self.ui["panel_2"])])

        style.configure("Primary.TButton", padding=(14, 10), background=self.ui["primary"], foreground="white")
        style.map("Primary.TButton", background=[("active", self.ui["primary_hover"])])

        style.configure("Secondary.TButton", padding=(14, 10), background=self.ui["panel_2"], foreground=self.ui["text"])
        style.map("Secondary.TButton", background=[("active", self.ui["panel"])])

        style.configure("Danger.TButton", padding=(14, 10), background=self.ui["danger"], foreground="white")
        style.map("Danger.TButton", background=[("active", "#DC2626")])

        style.configure("Chip.TLabel", background=self.ui["panel"], foreground=self.ui["muted"], padding=(8, 4))

    def _build_ui(self):
        root = ttk.Frame(self.root, style="App.TFrame", padding=(16, 14, 16, 16))
        root.pack(fill="both", expand=True)

        header = ttk.Frame(root, style="App.TFrame")
        header.pack(fill="x")
        ttk.Label(header, text="Voice Similarity & RVC Extractor", style="Title.TLabel").pack(side="left")
        ttk.Label(header, text="Professional acoustic comparison & report export", style="Sub.TLabel").pack(
            side="left", padx=(12, 0)
        )

        top = ttk.Frame(root, style="App.TFrame")
        top.pack(fill="x", pady=(14, 10))

        cards = ttk.Frame(top, style="App.TFrame")
        cards.pack(fill="x")
        cards.columnconfigure(0, weight=1)
        cards.columnconfigure(1, weight=1)

        self._build_audio_card(cards, col=0, kind="src")
        self._build_audio_card(cards, col=1, kind="tgt")

        actions = ttk.Frame(root, style="App.TFrame")
        actions.pack(fill="x", pady=(2, 10))

        ttk.Button(actions, text="Analyze", command=self.analyze, style="Primary.TButton").pack(side="left")
        ttk.Button(actions, text="Full Report", command=self.generate_full_report, style="Secondary.TButton").pack(
            side="left", padx=(10, 0)
        )

        self.status_chip = ttk.Label(actions, text="Idle", style="Chip.TLabel")
        self.status_chip.pack(side="right")

        out = ttk.Frame(root, style="Panel.TFrame", padding=(12, 10))
        out.pack(fill="both", expand=True)

        ttk.Label(out, text="Output", style="Value.TLabel").pack(anchor="w")
        body = ttk.Frame(out, style="Panel.TFrame")
        body.pack(fill="both", expand=True, pady=(8, 0))
        body.columnconfigure(0, weight=1)
        body.rowconfigure(0, weight=1)

        self.result = tk.Text(
            body,
            height=14,
            state="disabled",
            relief="flat",
            bd=0,
            highlightthickness=1,
            highlightbackground=self.ui["border"],
            highlightcolor=self.ui["focus"],
            bg=self.ui["bg"],
            fg=self.ui["text"],
            insertbackground=self.ui["text"],
            font=self.ui["mono"],
            wrap="word",
        )
        self.result.grid(row=0, column=0, sticky="nsew")
        yscroll = ttk.Scrollbar(body, orient="vertical", command=self.result.yview)
        yscroll.grid(row=0, column=1, sticky="ns")
        self.result.configure(yscrollcommand=yscroll.set)

    def _build_audio_card(self, parent, col, kind):
        title = "Source (your voice)" if kind == "src" else "Target (voice to clone)"
        icon = "🎙" if kind == "src" else "🧬"

        card = ttk.Frame(parent, style="Card.TFrame", padding=(14, 12))
        card.grid(row=0, column=col, sticky="nsew", padx=(0 if col == 0 else 10), pady=0)

        head = ttk.Frame(card, style="Card.TFrame")
        head.pack(fill="x")
        ttk.Label(head, text=f"{icon}  {title}", style="Value.TLabel").pack(side="left")

        row = ttk.Frame(card, style="Card.TFrame")
        row.pack(fill="x", pady=(12, 0))
        ttk.Label(row, text="ID", style="Field.TLabel").pack(side="left")

        entry = ttk.Entry(row, width=10, style="Modern.TEntry")
        entry.pack(side="left", padx=(8, 10))

        load_btn = ttk.Button(
            row, text="Load", command=self.load_src if kind == "src" else self.load_tgt, style="Modern.TButton"
        )
        load_btn.pack(side="left")

        rec_btn = ttk.Button(
            row, text="Record", command=self.toggle_src if kind == "src" else self.toggle_tgt, style="Modern.TButton"
        )
        rec_btn.pack(side="left", padx=(8, 0))

        play_btn = ttk.Button(
            row,
            text="Play",
            command=(lambda: self.play(self.src_audio, self.src_sr))
            if kind == "src"
            else (lambda: self.play(self.tgt_audio, self.tgt_sr)),
            style="Modern.TButton",
        )
        play_btn.pack(side="left", padx=(8, 0))

        status = ttk.Label(card, text="Not loaded", style="Field.TLabel")
        status.pack(anchor="w", pady=(10, 0))

        if kind == "src":
            self.src_id = entry
            self.btn_src_rec = rec_btn
            self.lbl_src = status
        else:
            self.tgt_id = entry
            self.btn_tgt_rec = rec_btn
            self.lbl_tgt = status

    def _set_status(self, text, tone="muted"):
        if not hasattr(self, "status_chip"):
            return
        colors = {
            "muted": self.ui["muted"],
            "ok": self.ui["success"],
            "warn": self.ui["warn"],
            "bad": self.ui["danger"],
        }
        self.status_chip.configure(text=text, foreground=colors.get(tone, self.ui["muted"]))

    def load_src(self):
        path = filedialog.askopenfilename(filetypes=[("Audio files", "*.wav *.mp3 *.flac")])
        if not path: return
        try:
            self._set_status("Loading source…", "muted")
            y, sr = sf.read(path, dtype='float64')
            if y.ndim > 1: y = y[:, 0]
            self.src_audio = y; self.src_sr = sr
            self.lbl_src.configure(text=f"Loaded ({len(y)/sr:.1f}s)", foreground=self.ui["success"])
            self._set_status("Source ready", "ok")
        except Exception as e:
            self._set_status("Source load failed", "bad")
            messagebox.showerror("Error", str(e))

    def load_tgt(self):
        path = filedialog.askopenfilename(filetypes=[("Audio files", "*.wav *.mp3 *.flac")])
        if not path: return
        try:
            self._set_status("Loading target…", "muted")
            y, sr = sf.read(path, dtype='float64')
            if y.ndim > 1: y = y[:, 0]
            self.tgt_audio = y; self.tgt_sr = sr
            self.lbl_tgt.configure(text=f"Loaded ({len(y)/sr:.1f}s)", foreground=self.ui["success"])
            self._set_status("Target ready", "ok")
        except Exception as e:
            self._set_status("Target load failed", "bad")
            messagebox.showerror("Error", str(e))

    def toggle_src(self):
        if not self.src_recorder.recording:
            if self.tgt_recorder.recording:
                messagebox.showwarning("Conflict", "Stop target recording first.")
                return
            self.src_recorder.start()
            self.btn_src_rec.configure(text="Stop", style="Danger.TButton")
            self.lbl_src.configure(text="Recording…", foreground=self.ui["focus"])
            self._set_status("Recording source…", "warn")
        else:
            audio = self.src_recorder.stop()
            self.btn_src_rec.configure(text="Record", style="Modern.TButton")
            if audio is not None:
                self.src_audio = audio; self.src_sr = 16000
                self.lbl_src.configure(text=f"Recorded ({len(audio)/16000:.1f}s)", foreground=self.ui["success"])
                self._set_status("Source ready", "ok")
            else:
                self._set_status("Recording stopped", "muted")

    def toggle_tgt(self):
        if not self.tgt_recorder.recording:
            if self.src_recorder.recording:
                messagebox.showwarning("Conflict", "Stop source recording first.")
                return
            self.tgt_recorder.start()
            self.btn_tgt_rec.configure(text="Stop", style="Danger.TButton")
            self.lbl_tgt.configure(text="Recording…", foreground=self.ui["focus"])
            self._set_status("Recording target…", "warn")
        else:
            audio = self.tgt_recorder.stop()
            self.btn_tgt_rec.configure(text="Record", style="Modern.TButton")
            if audio is not None:
                self.tgt_audio = audio; self.tgt_sr = 16000
                self.lbl_tgt.configure(text=f"Recorded ({len(audio)/16000:.1f}s)", foreground=self.ui["success"])
                self._set_status("Target ready", "ok")
            else:
                self._set_status("Recording stopped", "muted")

    def play(self, audio, sr):
        if audio is None: messagebox.showwarning("Warning", "No audio"); return
        sd.stop(); sd.play(audio, samplerate=sr)

    def analyze(self):
        if self.src_audio is None or self.tgt_audio is None:
            messagebox.showwarning("Warning", "Load both audios first")
            return
        try:
            self._set_status("Analyzing…", "muted")
            decision = analyze_speaker_verification(self.src_audio, self.src_sr, self.tgt_audio, self.tgt_sr)
            src_feat = get_audio_features(self.src_audio, self.src_sr)
            tgt_feat = get_audio_features(self.tgt_audio, self.tgt_sr)
            # fix #1: explicit 3-tuple unpacking (was fragile `[1:]` slice)
            _, _, src_gender = extract_pitch_array(self.src_audio, self.src_sr)
            _, _, tgt_gender = extract_pitch_array(self.tgt_audio, self.tgt_sr)

            res = (
                f"Source (ID: {self.src_id.get().strip() or 'None'})\n"
                f"  Pitch: {decision['src_pitch']:.2f} Hz  Gender: {src_gender:.2f}  Duration: {src_feat['duration']}s\n"
                f"  RMS: {src_feat['rms']}  ZCR: {src_feat['zcr']}  Centroid: {src_feat['centroid']} Hz\n\n"
                f"Target (ID: {self.tgt_id.get().strip() or 'None'})\n"
                f"  Pitch: {decision['tgt_pitch']:.2f} Hz  Gender: {tgt_gender:.2f}  Duration: {tgt_feat['duration']}s\n"
                f"  RMS: {tgt_feat['rms']}  ZCR: {tgt_feat['zcr']}  Centroid: {tgt_feat['centroid']} Hz\n\n"
                f"--- Multi‑Metric Speaker Verification ---\n"
                f"  Speaker Embedding: {decision['speaker_sim']:.3f}  (timbre)\n"
                f"  Pitch Similarity:  {decision['pitch_sim']:.3f}  (shift {decision['semitones']:.1f} st)\n"
                f"  Formant F2 diff:   {decision['f2_diff']:.0f} Hz  (norm {decision['f2_norm']:.2f})\n"
                f"  Spectral Centroid diff: {decision['cent_diff']:.0f} Hz  (norm {decision['cent_norm']:.2f})\n"
                f"  MFCC Similarity:   {decision['mfcc_sim']:.3f}\n"
                f"  Decision Score:    {decision['decision_score']:.3f}\n\n"
                f"  Result: {decision['verdict']} (Confidence: {decision['confidence']})\n"
                f"  Reason: {decision['reason']}"
            )
            self.result.config(state="normal")
            self.result.delete(1.0, tk.END)
            self.result.insert(tk.END, res)
            self.result.config(state="disabled")
            tone = "ok" if decision.get("verdict") == "YES" else ("warn" if decision.get("verdict") == "Uncertain" else "bad")
            self._set_status(f"Result: {decision.get('verdict', 'Done')}", tone)
        except Exception as e:
            self._set_status("Analysis failed", "bad")
            messagebox.showerror("Analysis Error", str(e))

    def generate_full_report(self):
        if self.src_audio is None or self.tgt_audio is None:
            messagebox.showwarning("Warning", "Load both audios first")
            return
        try:
            self._set_status("Generating report…", "muted")
            src_id = self.src_id.get().strip() or "src"
            tgt_id = self.tgt_id.get().strip() or "tgt"

            decision = analyze_speaker_verification(self.src_audio, self.src_sr, self.tgt_audio, self.tgt_sr)
            src_dur = len(self.src_audio) / self.src_sr
            tgt_dur = len(self.tgt_audio) / self.tgt_sr
            common_dur = min(src_dur, tgt_dur)

            metrics = {
                "Source Pitch (Hz)": f"{decision['src_pitch']:.2f}",
                "Target Pitch (Hz)": f"{decision['tgt_pitch']:.2f}",
                "Semitone Shift": f"{decision['semitones']:.1f}",
                "Speaker Embedding": f"{decision['speaker_sim']:.3f}",
                "Pitch Similarity": f"{decision['pitch_sim']:.3f}",
                "F2 Difference (Hz)": f"{decision['f2_diff']:.0f}",
                "F2 Normalized": f"{decision['f2_norm']:.2f}",
                "Centroid Diff (Hz)": f"{decision['cent_diff']:.0f}",
                "MFCC Similarity": f"{decision['mfcc_sim']:.3f}",
                "Decision Score": f"{decision['decision_score']:.3f}",
                "Verdict": f"{decision['verdict']} ({decision['confidence']})",
                "Aligned Duration (s)": f"{common_dur:.1f}"
            }
            metrics = {k: str(v) for k, v in metrics.items()}

            # Save figures to picture directory (unique name with IDs)
            picture_subdir = os.path.join(self.picture_dir, f"{src_id}_{tgt_id}")
            img_paths = generate_report_figures(
                self.src_audio, self.src_sr, self.tgt_audio, self.tgt_sr,
                src_id, tgt_id, metrics, decision,
                output_dir=picture_subdir
            )

            # Generate Word report in report directory
            if DOCX_AVAILABLE:
                docx_path = os.path.join(self.report_dir, f"report_{src_id}_{tgt_id}.docx")
                success = create_word_report(img_paths, src_id, tgt_id, metrics, decision, docx_path)
                if success:
                    open_file_with_default_app(docx_path)
                    msg = f"Word report saved to:\n{docx_path}\nOpening now..."
                else:
                    msg = "Word report creation failed."
            else:
                # Fallback: combine images into a big PNG in report dir
                combined_path = os.path.join(self.report_dir, f"report_{src_id}_{tgt_id}.png")
                if combine_images_vertically(img_paths, combined_path):
                    open_file_with_default_app(combined_path)
                    msg = f"Combined image saved to:\n{combined_path}\nOpening now..."
                else:
                    open_file_with_default_app(img_paths[0])
                    msg = f"Opening first plot:\n{img_paths[0]}"

            # Save text summary in report directory
            txt_path = os.path.join(self.report_dir, f"summary_{src_id}_{tgt_id}.txt")
            with open(txt_path, 'w', encoding='utf-8') as f:
                f.write(f"Voice Similarity Report: {src_id} vs {tgt_id}\n")
                for k, v in metrics.items():
                    f.write(f"{k}: {v}\n")
                f.write(f"\nVerdict: {decision['verdict']} ({decision['confidence']})\n")
                f.write(f"Reason: {decision['reason']}\n")

            summary = (
                f"Report generated for {src_id} vs {tgt_id}\n"
                f"Verdict: {decision['verdict']} ({decision['confidence']})\n"
                f"{decision['reason']}\n\n{msg}"
            )
            self.result.config(state="normal")
            self.result.delete(1.0, tk.END)
            self.result.insert(tk.END, summary)
            self.result.config(state="disabled")
            self._set_status("Report ready", "ok")
        except Exception as e:
            self._set_status("Report failed", "bad")
            messagebox.showerror("Report Error", str(e))

if __name__ == "__main__":
    root = tk.Tk()
    app = PitchGUI(root)
    root.mainloop()
