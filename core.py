#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
core.py — Acoustic analysis core for RVC_factor.

Extracted from legacy/factor.py with the following fixes applied:
  #1  Tuple unpacking now explicit (no more `[1:]` slicing trick)
  #2  F0 extraction parameters unified (PITCH_FLOOR / PITCH_CEILING constants)
  #3  Bare `except:` replaced by `except Exception:`
  #4  File-open uses subprocess with arg list (no shell injection)
  #5  VoiceEncoder cached as a module-level singleton
  #6  MFCC computed once per analyze flow, not twice

This module has NO GUI / NO Flask code, so it can be imported by both the
Tkinter desktop backup (legacy/factor.py) and the web app (app.py).
"""

import os
import sys
import subprocess
import warnings
import numpy as np
import parselmouth
import librosa
import librosa.display

# ---------- Optional imports ----------
try:
    from resemblyzer import VoiceEncoder, preprocess_wav
    RESEMBLYZER_AVAILABLE = True
except ImportError:
    RESEMBLYZER_AVAILABLE = False

try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

# Report generation needs matplotlib + docx; imported lazily where used.

warnings.filterwarnings("ignore")

# ==================== Constants (fix #2: unify F0 params) ====================
PITCH_TIME_STEP = 0.01
PITCH_FLOOR = 65       # Hz — wide range, suits male-low to high-pitched voices
PITCH_CEILING = 2093   # Hz
FORMANT_TIME_STEP = 0.01
FORMANT_MAX_NUMBER = 5
FORMANT_MAXIMUM = 5500
FORMANT_WINDOW_LENGTH = 0.025

# Gender mapping anchors (median pitch in Hz)
GENDER_LOW = 120
GENDER_HIGH = 210

# ==================== Voice encoder singleton (fix #5) ====================
_ENCODER = None


def get_voice_encoder():
    """Return a cached VoiceEncoder, loading it once on first use."""
    global _ENCODER
    if not RESEMBLYZER_AVAILABLE:
        return None
    if _ENCODER is None:
        _ENCODER = VoiceEncoder()
    return _ENCODER


# ==================== Audio feature extraction ====================
def extract_pitch_array(audio, sr=16000):
    """Return (f0_array, median_pitch_hz, gender_0_to_1)."""
    snd = parselmouth.Sound(audio, sampling_frequency=sr)
    pitch = snd.to_pitch(time_step=PITCH_TIME_STEP,
                         pitch_floor=PITCH_FLOOR,
                         pitch_ceiling=PITCH_CEILING)
    f0 = pitch.selected_array['frequency']
    valid = f0[f0 > 0]
    if len(valid) == 0:
        raise ValueError("No voiced frames detected")
    median_pitch = np.median(valid)
    gender = np.clip((median_pitch - GENDER_LOW) / (GENDER_HIGH - GENDER_LOW), 0.0, 1.0)
    return f0, median_pitch, gender


def to_pitch_obj(audio, sr):
    """Build a Praat Pitch object using the unified F0 parameters."""
    snd = parselmouth.Sound(audio, sampling_frequency=sr)
    return snd.to_pitch(time_step=PITCH_TIME_STEP,
                        pitch_floor=PITCH_FLOOR,
                        pitch_ceiling=PITCH_CEILING)


def extract_formants(audio, sr=16000):
    """Return (times, f1, f2, f3) as numpy arrays."""
    snd = parselmouth.Sound(audio, sampling_frequency=sr)
    formant = snd.to_formant_burg(time_step=FORMANT_TIME_STEP,
                                  max_number_of_formants=FORMANT_MAX_NUMBER,
                                  maximum_formant=FORMANT_MAXIMUM,
                                  window_length=FORMANT_WINDOW_LENGTH)
    times = formant.xs()
    f1 = [formant.get_value_at_time(1, t) for t in times]
    f2 = [formant.get_value_at_time(2, t) for t in times]
    f3 = [formant.get_value_at_time(3, t) for t in times]
    return np.array(times), np.array(f1), np.array(f2), np.array(f3)


def get_audio_features(audio, sr):
    """Return basic scalar features for an audio clip."""
    dur = len(audio) / sr
    rms = np.sqrt(np.mean(audio ** 2))
    zcr = np.mean(librosa.feature.zero_crossing_rate(y=audio))
    cent = np.mean(librosa.feature.spectral_centroid(y=audio, sr=sr))
    return {
        'duration': round(dur, 2),
        'rms': round(float(rms), 4),
        'zcr': round(float(zcr), 4),
        'centroid': round(float(cent), 1),
    }


def analyze_single_audio(audio, sr):
    """Return single-clip acoustic metrics computed from the raw waveform."""
    f0, pitch, gender = extract_pitch_array(audio, sr)
    t_f, f1, f2, f3 = extract_formants(audio, sr)
    features = get_audio_features(audio, sr)
    mfcc = librosa.feature.mfcc(y=audio.astype(np.float64), sr=sr, n_mfcc=13)
    bandwidth = librosa.feature.spectral_bandwidth(y=audio.astype(np.float64), sr=sr)
    rolloff = librosa.feature.spectral_rolloff(y=audio.astype(np.float64), sr=sr)
    valid_f0 = f0[f0 > 0]
    return {
        **features,
        "pitch_hz": float(pitch),
        "gender": float(gender),
        "pitch_low": float(np.percentile(valid_f0, 10)),
        "pitch_high": float(np.percentile(valid_f0, 90)),
        "pitch_std": float(np.std(valid_f0)),
        "f1_mean": float(np.nanmean(f1)),
        "f2_mean": float(np.nanmean(f2)),
        "f3_mean": float(np.nanmean(f3)),
        "spectral_bandwidth": float(np.nanmean(bandwidth)),
        "spectral_rolloff": float(np.nanmean(rolloff)),
        "mfcc_texture": float(np.mean(np.std(mfcc, axis=1))),
    }


# ==================== Similarity metrics ====================
def compute_pitch_similarity(p1, p2):
    semis = abs(12 * np.log2(p2 / p1))
    return float(np.clip(1 - semis / 12.0, 0.0, 1.0))


def compute_mfcc(a1, sr1, a2, sr2):
    """Return raw cosine similarity in [-1, 1] between MFCC means, plus the
    shifted-to-[0,1] value used by callers."""
    m1 = librosa.feature.mfcc(y=a1.astype(np.float64), sr=sr1, n_mfcc=13)
    m2 = librosa.feature.mfcc(y=a2.astype(np.float64), sr=sr2, n_mfcc=13)
    v1, v2 = np.mean(m1, axis=1), np.mean(m2, axis=1)
    dot = np.dot(v1, v2)
    norm = np.linalg.norm(v1) * np.linalg.norm(v2)
    raw = float(np.clip(dot / norm, -1, 1)) if norm > 0 else 0.0
    return raw, float(np.clip((raw + 1) / 2.0, 0.0, 1.0))


def compute_speaker_similarity(a1, sr1, a2, sr2, mfcc_raw=None):
    """Speaker-embedding similarity (Resemblyzer) when available, else MFCC fallback.

    Caller may pass a precomputed MFCC raw cosine to avoid recomputation (fix #6).
    """
    if RESEMBLYZER_AVAILABLE:
        try:
            enc = get_voice_encoder()
            w1 = preprocess_wav(a1.astype(np.float64), source_sr=sr1)
            w2 = preprocess_wav(a2.astype(np.float64), source_sr=sr2)
            e1, e2 = enc.embed_utterance(w1), enc.embed_utterance(w2)
            return float(np.clip((np.dot(e1, e2) + 1) / 2.0, 0.0, 1.0))
        except Exception as e:  # fix #3: never swallow KeyboardInterrupt
            print(f"Resemblyzer failed: {e}")
    if mfcc_raw is None:
        mfcc_raw, _ = compute_mfcc(a1, sr1, a2, sr2)
    return float(np.clip((mfcc_raw + 1) / 2.0, 0.0, 1.0))


# ==================== Multi-metric decision engine ====================
def analyze_speaker_verification(src_audio, src_sr, tgt_audio, tgt_sr):
    """Return a dict with all metrics + a final verdict and human-readable reason.

    Computes every expensive quantity (MFCC, embedding, pitch, formants) exactly
    once and reuses it across the score, the reasons, and the figure generator.
    """
    # 1. MFCC cosine (raw + [0,1]) — compute once, reuse below (fix #6)
    mfcc_raw, mfcc_sim_01 = compute_mfcc(src_audio, src_sr, tgt_audio, tgt_sr)

    # 2. Speaker embedding (uses cached encoder; reuses MFCC on fallback)
    spk_sim = compute_speaker_similarity(src_audio, src_sr, tgt_audio, tgt_sr, mfcc_raw=mfcc_raw)

    # 3. Pitch
    _, src_pitch, _ = extract_pitch_array(src_audio, src_sr)
    _, tgt_pitch, _ = extract_pitch_array(tgt_audio, tgt_sr)
    pitch_sim = compute_pitch_similarity(src_pitch, tgt_pitch)
    semitones = 12 * np.log2(tgt_pitch / src_pitch)

    # 4. Formant F2 means
    _, _, f2_src, _ = extract_formants(src_audio, src_sr)
    _, _, f2_tgt, _ = extract_formants(tgt_audio, tgt_sr)
    f2_diff = abs(float(np.nanmean(f2_src)) - float(np.nanmean(f2_tgt)))
    f2_norm = float(np.clip(f2_diff / 300.0, 0.0, 1.0))  # 300 Hz reference

    # 5. Spectral centroid
    src_feat = get_audio_features(src_audio, src_sr)
    tgt_feat = get_audio_features(tgt_audio, tgt_sr)
    cent_diff = abs(src_feat['centroid'] - tgt_feat['centroid'])
    cent_norm = float(np.clip(cent_diff / 1000.0, 0.0, 1.0))

    # Weighted decision score (industry-inspired)
    decision_score = (0.4 * spk_sim +
                      0.2 * pitch_sim +
                      0.2 * (1 - f2_norm) +
                      0.1 * (1 - cent_norm) +
                      0.1 * mfcc_sim_01)

    if spk_sim > 0.85 and decision_score > 0.78:
        verdict, confidence = "YES", "High"
    elif spk_sim > 0.78 and decision_score > 0.68 and pitch_sim > 0.6 and f2_norm < 0.5:
        verdict, confidence = "YES", "Medium"
    elif spk_sim > 0.70 and decision_score > 0.60 and pitch_sim > 0.5:
        verdict, confidence = "Uncertain", "Low"
    else:
        verdict, confidence = "NO", "Very Low / Different"

    reasons = []
    if spk_sim > 0.82:
        reasons.append("说话人音色特征非常接近")
    elif spk_sim > 0.70:
        reasons.append("说话人音色特征有一定相似")
    else:
        reasons.append("说话人音色特征差异明显")

    if pitch_sim > 0.75:
        reasons.append("音高范围很接近")
    elif pitch_sim > 0.5:
        reasons.append(f"音高相差约 {abs(semitones):.1f} 个半音，属于中等差异")
    else:
        reasons.append(f"音高相差约 {abs(semitones):.1f} 个半音，差异较大")

    if f2_norm < 0.3:
        reasons.append("F2 共振峰均值很接近，提示口腔共鸣位置相似")
    elif f2_norm < 0.6:
        reasons.append(f"F2 共振峰相差约 {f2_diff:.0f} Hz，有可见差异")
    else:
        reasons.append(f"F2 共振峰相差约 {f2_diff:.0f} Hz，差异明显")

    if cent_norm < 0.3:
        reasons.append("频谱质心很接近，声音明亮度相似")
    elif cent_norm < 0.6:
        reasons.append("频谱质心有中等差异，声音明亮度不完全一致")
    else:
        reasons.append("频谱质心差异明显，声音明亮度不同")

    full_reason = "；".join(reasons) + "。"
    if verdict == "YES":
        full_reason += "多个声音线索方向一致，因此支持同一说话人的判断。"
    elif verdict == "NO":
        full_reason += "多个声音线索分歧较大，因此更支持不同说话人的判断。"
    else:
        full_reason += "目前线索有相似也有差异，建议结合人工听辨和录音条件再复核。"

    return {
        'speaker_sim': spk_sim,
        'pitch_sim': pitch_sim,
        'f2_diff': f2_diff,
        'f2_norm': f2_norm,
        'cent_diff': cent_diff,
        'cent_norm': cent_norm,
        'mfcc_sim': mfcc_sim_01,
        'decision_score': decision_score,
        'verdict': verdict,
        'confidence': confidence,
        'reason': full_reason,
        'src_pitch': float(src_pitch),
        'tgt_pitch': float(tgt_pitch),
        'semitones': float(semitones),
        'src_centroid': src_feat['centroid'],
        'tgt_centroid': tgt_feat['centroid'],
        'src_features': src_feat,
        'tgt_features': tgt_feat,
    }


# ==================== Plotting ====================
def _annotate(ax, text, x=0.02, y=0.95, fontsize=8):
    ax.text(x, y, str(text), transform=ax.transAxes, fontsize=fontsize,
            verticalalignment='top',
            bbox=dict(boxstyle='round', facecolor='white', alpha=0.85))


def _setup_matplotlib_cjk(matplotlib):
    """Register a local CJK font so report PNGs can render Chinese labels."""
    from matplotlib import font_manager

    font_candidates = [
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\msyhbd.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
    ]
    for font_path in font_candidates:
        if os.path.exists(font_path):
            font_manager.fontManager.addfont(font_path)
            family = font_manager.FontProperties(fname=font_path).get_name()
            matplotlib.rcParams['font.family'] = family
            matplotlib.rcParams['font.sans-serif'] = [family, 'DejaVu Sans']
            matplotlib.rcParams['axes.unicode_minus'] = False
            return True
    matplotlib.rcParams['font.sans-serif'] = ['DejaVu Sans']
    matplotlib.rcParams['axes.unicode_minus'] = False
    return False


def generate_report_figures(src_audio, src_sr, tgt_audio, tgt_sr,
                            src_id, tgt_id, output_dir):
    """Save the 5 standard report figures to output_dir. Returns list of paths."""
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    _setup_matplotlib_cjk(matplotlib)

    os.makedirs(output_dir, exist_ok=True)
    img_paths = []
    src_label = str(src_id)
    tgt_label = str(tgt_id)

    src_dur = len(src_audio) / src_sr
    tgt_dur = len(tgt_audio) / tgt_sr
    common_dur = min(src_dur, tgt_dur)
    src_aligned = src_audio[:int(common_dur * src_sr)] if src_dur > common_dur else src_audio
    tgt_aligned = tgt_audio[:int(common_dur * tgt_sr)] if tgt_dur > common_dur else tgt_audio

    src_f0, src_pitch, _ = extract_pitch_array(src_audio, src_sr)
    tgt_f0, tgt_pitch, _ = extract_pitch_array(tgt_audio, tgt_sr)

    # ---- Fig 1: Overview ----
    fig1, axes = plt.subplots(3, 2, figsize=(10, 7))
    fig1.suptitle(f"声音总览：{src_label} vs {tgt_label}（前 {common_dur:.1f} 秒）", fontweight='bold')
    t_src = np.linspace(0, len(src_aligned) / src_sr, len(src_aligned))
    t_tgt = np.linspace(0, len(tgt_aligned) / tgt_sr, len(tgt_aligned))
    axes[0, 0].plot(t_src, src_aligned, color='steelblue', linewidth=0.7)
    axes[0, 0].set_title(f"{src_label} - 波形", fontsize=10)
    axes[0, 1].plot(t_tgt, tgt_aligned, color='darkorange', linewidth=0.7)
    axes[0, 1].set_title(f"{tgt_label} - 波形", fontsize=10)
    for ax, audio, sr, title in [(axes[1, 0], src_aligned, src_sr, f"{src_label} - 频谱图"),
                                 (axes[1, 1], tgt_aligned, tgt_sr, f"{tgt_label} - 频谱图")]:
        D = librosa.stft(audio.astype(np.float64))
        S_db = librosa.amplitude_to_db(np.abs(D), ref=np.max)
        librosa.display.specshow(S_db, sr=sr, x_axis='time', y_axis='hz', ax=ax, cmap='inferno')
        ax.set_title(title, fontsize=10)
        ax.set_ylim(0, 3000)
    axes[2, 0].hist(src_f0[src_f0 > 0], bins=30, color='steelblue', alpha=0.7)
    axes[2, 0].axvline(src_pitch, color='red', linestyle='--', label=f"中位数 {src_pitch:.1f} Hz")
    axes[2, 0].set_title(f"{src_label} - 基频分布", fontsize=10)
    axes[2, 0].legend(fontsize=8)
    axes[2, 1].hist(tgt_f0[tgt_f0 > 0], bins=30, color='darkorange', alpha=0.7)
    axes[2, 1].axvline(tgt_pitch, color='red', linestyle='--', label=f"中位数 {tgt_pitch:.1f} Hz")
    axes[2, 1].set_title(f"{tgt_label} - 基频分布", fontsize=10)
    axes[2, 1].legend(fontsize=8)
    plt.tight_layout()
    p1 = os.path.join(output_dir, "overview.png"); fig1.savefig(p1, dpi=150); plt.close(fig1)
    img_paths.append(p1)

    # ---- Fig 2: F0 contours ----
    pitch_src = to_pitch_obj(src_aligned, src_sr)
    pitch_tgt = to_pitch_obj(tgt_aligned, tgt_sr)
    fig2, (ax1, ax2) = plt.subplots(2, 1, figsize=(8, 4))
    ax1.plot(pitch_src.xs(), pitch_src.selected_array['frequency'], color='steelblue', linewidth=1)
    ax1.set_title(f"{src_label} - 音高走势", fontsize=10)
    _annotate(ax1, f"中位数：{src_pitch:.1f} Hz")
    ax2.plot(pitch_tgt.xs(), pitch_tgt.selected_array['frequency'], color='darkorange', linewidth=1)
    ax2.set_title(f"{tgt_label} - 音高走势", fontsize=10)
    _annotate(ax2, f"中位数：{tgt_pitch:.1f} Hz")
    plt.tight_layout()
    p2 = os.path.join(output_dir, "f0_contours.png"); fig2.savefig(p2, dpi=150); plt.close(fig2)
    img_paths.append(p2)

    # ---- Fig 3: Formants ----
    t_s, f1_s, f2_s, f3_s = extract_formants(src_aligned, src_sr)
    t_t, f1_t, f2_t, f3_t = extract_formants(tgt_aligned, tgt_sr)
    fig3, (ax1, ax2) = plt.subplots(2, 1, figsize=(8, 4))
    ax1.plot(t_s, f1_s, 'r', label='F1'); ax1.plot(t_s, f2_s, 'g', label='F2'); ax1.plot(t_s, f3_s, 'b', label='F3')
    ax1.set_title(f"{src_label} - 共振峰", fontsize=10); ax1.legend(fontsize=8)
    _annotate(ax1, f"平均 F1={np.nanmean(f1_s):.0f} Hz，F2={np.nanmean(f2_s):.0f} Hz")
    ax2.plot(t_t, f1_t, 'r', label='F1'); ax2.plot(t_t, f2_t, 'g', label='F2'); ax2.plot(t_t, f3_t, 'b', label='F3')
    ax2.set_title(f"{tgt_label} - 共振峰", fontsize=10); ax2.legend(fontsize=8)
    _annotate(ax2, f"平均 F1={np.nanmean(f1_t):.0f} Hz，F2={np.nanmean(f2_t):.0f} Hz")
    plt.tight_layout()
    p3 = os.path.join(output_dir, "formants.png"); fig3.savefig(p3, dpi=150); plt.close(fig3)
    img_paths.append(p3)

    # ---- Fig 4: Average spectrum ----
    fig4, ax = plt.subplots(figsize=(6, 4))
    for audio, sr, label, color in [(src_audio, src_sr, src_label, 'steelblue'),
                                    (tgt_audio, tgt_sr, tgt_label, 'darkorange')]:
        D = np.abs(librosa.stft(audio.astype(np.float64)))
        avg = np.mean(D, axis=1)
        freqs = librosa.fft_frequencies(sr=sr)
        db = librosa.amplitude_to_db(avg, ref=np.max)
        ax.semilogx(freqs, db, color=color, label=label, linewidth=1)
    ax.set_xlim(50, max(src_sr, tgt_sr) // 2)
    ax.set_xlabel("频率 (Hz)"); ax.set_ylabel("dB")
    ax.set_title("平均频谱（完整音频）")
    ax.legend(fontsize=8); ax.grid(True, alpha=0.3)
    _annotate(ax, f"频谱质心：咨询师={get_audio_features(src_audio, src_sr)['centroid']} Hz，"
                 f"来访者={get_audio_features(tgt_audio, tgt_sr)['centroid']} Hz")
    plt.tight_layout()
    p4 = os.path.join(output_dir, "avg_spectrum.png"); fig4.savefig(p4, dpi=150); plt.close(fig4)
    img_paths.append(p4)

    # ---- Fig 5: MFCC ----
    fig5, (ax1, ax2) = plt.subplots(2, 1, figsize=(6, 5))
    mfcc_s = librosa.feature.mfcc(y=src_aligned.astype(np.float64), sr=src_sr, n_mfcc=13)
    mfcc_t = librosa.feature.mfcc(y=tgt_aligned.astype(np.float64), sr=tgt_sr, n_mfcc=13)
    librosa.display.specshow(mfcc_s, sr=src_sr, x_axis='time', ax=ax1, cmap='coolwarm')
    ax1.set_title(f"{src_label} - 声音纹理（MFCC）", fontsize=10)
    librosa.display.specshow(mfcc_t, sr=tgt_sr, x_axis='time', ax=ax2, cmap='coolwarm')
    ax2.set_title(f"{tgt_label} - 声音纹理（MFCC）", fontsize=10)
    plt.tight_layout()
    p5 = os.path.join(output_dir, "mfcc.png"); fig5.savefig(p5, dpi=150); plt.close(fig5)
    img_paths.append(p5)

    return img_paths


def generate_single_report_figures(audio, sr, label, output_dir):
    """Save the standard single-audio report figures to output_dir."""
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    _setup_matplotlib_cjk(matplotlib)

    os.makedirs(output_dir, exist_ok=True)
    img_paths = []
    label = str(label)
    duration = len(audio) / sr
    t = np.linspace(0, duration, len(audio))
    f0, pitch, _ = extract_pitch_array(audio, sr)
    pitch_obj = to_pitch_obj(audio, sr)
    formant_t, f1, f2, f3 = extract_formants(audio, sr)
    features = get_audio_features(audio, sr)

    # ---- Fig 1: Overview ----
    fig1, axes = plt.subplots(3, 1, figsize=(9, 7))
    fig1.suptitle(f"{label} - 单人声音总览", fontweight="bold")
    axes[0].plot(t, audio, color="steelblue", linewidth=0.7)
    axes[0].set_title("波形：看音量随时间的起伏", fontsize=10)
    axes[0].set_xlabel("时间 (秒)")
    D = librosa.stft(audio.astype(np.float64))
    S_db = librosa.amplitude_to_db(np.abs(D), ref=np.max)
    librosa.display.specshow(S_db, sr=sr, x_axis="time", y_axis="hz", ax=axes[1], cmap="inferno")
    axes[1].set_ylim(0, 3000)
    axes[1].set_title("频谱图：看不同频率的声音强弱", fontsize=10)
    valid_f0 = f0[f0 > 0]
    axes[2].hist(valid_f0, bins=30, color="steelblue", alpha=0.75)
    axes[2].axvline(pitch, color="red", linestyle="--", label=f"中位数 {pitch:.1f} Hz")
    axes[2].set_title("基频分布：看常用音高集中在哪里", fontsize=10)
    axes[2].legend(fontsize=8)
    plt.tight_layout()
    p1 = os.path.join(output_dir, "single_overview.png"); fig1.savefig(p1, dpi=150); plt.close(fig1)
    img_paths.append(p1)

    # ---- Fig 2: F0 contour ----
    fig2, ax = plt.subplots(figsize=(8, 4))
    ax.plot(pitch_obj.xs(), pitch_obj.selected_array["frequency"], color="steelblue", linewidth=1)
    ax.set_title(f"{label} - 音高走势", fontsize=10)
    ax.set_xlabel("时间 (秒)"); ax.set_ylabel("Hz")
    _annotate(ax, f"中位数：{pitch:.1f} Hz；常用范围：{np.percentile(valid_f0, 10):.1f}-{np.percentile(valid_f0, 90):.1f} Hz")
    plt.tight_layout()
    p2 = os.path.join(output_dir, "single_f0_contour.png"); fig2.savefig(p2, dpi=150); plt.close(fig2)
    img_paths.append(p2)

    # ---- Fig 3: Formants ----
    fig3, ax = plt.subplots(figsize=(8, 4))
    ax.plot(formant_t, f1, "r", label="F1")
    ax.plot(formant_t, f2, "g", label="F2")
    ax.plot(formant_t, f3, "b", label="F3")
    ax.set_title(f"{label} - 共振峰走势", fontsize=10)
    ax.set_xlabel("时间 (秒)"); ax.set_ylabel("Hz")
    ax.legend(fontsize=8)
    _annotate(ax, f"平均 F1={np.nanmean(f1):.0f} Hz，F2={np.nanmean(f2):.0f} Hz，F3={np.nanmean(f3):.0f} Hz")
    plt.tight_layout()
    p3 = os.path.join(output_dir, "single_formants.png"); fig3.savefig(p3, dpi=150); plt.close(fig3)
    img_paths.append(p3)

    # ---- Fig 4: Average spectrum ----
    fig4, ax = plt.subplots(figsize=(7, 4))
    avg = np.mean(np.abs(D), axis=1)
    freqs = librosa.fft_frequencies(sr=sr)
    db = librosa.amplitude_to_db(avg, ref=np.max)
    ax.semilogx(freqs, db, color="steelblue", linewidth=1)
    ax.set_xlim(50, sr // 2)
    ax.set_xlabel("频率 (Hz)"); ax.set_ylabel("dB")
    ax.set_title("平均频谱：看声音整体偏亮还是偏厚")
    ax.grid(True, alpha=0.3)
    _annotate(ax, f"频谱质心：{features['centroid']} Hz")
    plt.tight_layout()
    p4 = os.path.join(output_dir, "single_avg_spectrum.png"); fig4.savefig(p4, dpi=150); plt.close(fig4)
    img_paths.append(p4)

    # ---- Fig 5: MFCC ----
    fig5, ax = plt.subplots(figsize=(7, 4))
    mfcc = librosa.feature.mfcc(y=audio.astype(np.float64), sr=sr, n_mfcc=13)
    librosa.display.specshow(mfcc, sr=sr, x_axis="time", ax=ax, cmap="coolwarm")
    ax.set_title(f"{label} - 声音纹理（MFCC）", fontsize=10)
    plt.tight_layout()
    p5 = os.path.join(output_dir, "single_mfcc.png"); fig5.savefig(p5, dpi=150); plt.close(fig5)
    img_paths.append(p5)

    return img_paths


def combine_images_vertically(img_paths, output_path):
    if not PIL_AVAILABLE:
        return False
    images = [Image.open(p) for p in img_paths]
    widths, heights = zip(*(i.size for i in images))
    combined = Image.new('RGB', (max(widths), sum(heights)), 'white')
    y_offset = 0
    for im in images:
        combined.paste(im, (0, y_offset))
        y_offset += im.size[1]
    combined.save(output_path)
    return True


def create_word_report(img_paths, src_id, tgt_id, metrics, decision, output_docx):
    from docx import Document
    from docx.shared import Inches, Pt

    doc = Document()
    doc.styles['Normal'].font.size = Pt(10)
    verdict_map = {"YES": "更像同一说话人", "NO": "更像不同说话人", "Uncertain": "暂时无法确定"}
    confidence_map = {"High": "高", "Medium": "中等", "Low": "低", "Very Low / Different": "很低"}
    doc.add_heading(f'声音相似度报告：{src_id} vs {tgt_id}', 0)
    doc.add_paragraph(f"结论：{verdict_map.get(decision['verdict'], decision['verdict'])}（置信度：{confidence_map.get(decision['confidence'], decision['confidence'])}）",
                      style='Intense Quote')
    doc.add_paragraph(
        f"综合分为 {decision['decision_score']:.3f}，音色相似度为 {decision['speaker_sim']:.3f}，"
        f"音高相差约 {abs(decision['semitones']):.1f} 个半音。"
        "这些数字用于辅助理解两段声音的接近程度，不应单独作为医学诊断或身份结论。"
    )

    doc.add_heading('量化指标', level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = '指标'
    hdr[1].text = '数值'
    for key, val in metrics.items():
        row = table.add_row().cells
        row[0].text = key
        row[1].text = str(val)

    doc.add_heading('图表说明', level=2)
    explanations = [
        ("图 1：总览",
         "把音量起伏、频谱和音高分布放在一起看，帮助直观观察两段声音的节奏、高低和声音颜色是否接近。"),
        ("图 2：音高走势",
         "看声音高低如何随时间变化。线条越像，说明说话时升高和降低的习惯越接近。"),
        ("图 3：共振峰",
         "反映声音经过口腔和鼻腔后的共鸣位置，是辅助判断音色来源的重要线索。"),
        ("图 4：平均频谱",
         "看整体声音偏亮还是偏厚。差异较大时，可能是人不同，也可能是录音环境不同。"),
        ("图 5：声音纹理",
         "可以理解为声音细节的纹理图。颜色分布越接近，整体音色结构通常越相似。"),
    ]
    for i, (title, desc) in enumerate(explanations, 1):
        doc.add_heading(title, level=3)
        doc.add_paragraph(desc)
        if i - 1 < len(img_paths):
            doc.add_picture(img_paths[i - 1], width=Inches(5.5))
        doc.add_paragraph()

    doc.save(output_docx)
    return True


def open_file_with_default_app(filepath):
    """Open a file cross-platform using arg lists (no shell injection — fix #4)."""
    if sys.platform == "win32":
        os.startfile(filepath)
    elif sys.platform == "darwin":
        subprocess.run(["open", filepath], check=False)
    else:
        subprocess.run(["xdg-open", filepath], check=False)
